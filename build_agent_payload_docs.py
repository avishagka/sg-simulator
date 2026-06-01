from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCS = [
    ("02_screen_population_spec.md", "02_screen_population_spec.docx"),
    ("03_output_contract_universal_provider_payload.md", "Output Contract - UniversalProviderDisplayPayload.docx"),
    ("04_mapping_spec_universal_provider_payload.md", "Mapping Specification - UniversalProviderDisplayPayload.docx"),
    ("05_business_rules_universal_provider_payload.md", "Business Rules - UniversalProviderDisplayPayload.docx"),
]


def set_bidi(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_bidi(p)
    run = p.add_run(text.strip())
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "DADCE0")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def split_table_row(line):
    stripped = line.strip().strip("|")
    return [part.strip().replace("<br>", "\n") for part in stripped.split("|")]


def is_separator_row(line):
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = True
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(len(table.columns)):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, text, bold=(r_idx == 0))
            if r_idx == 0:
                shade_cell(cell, "E8EEF5")
    doc.add_paragraph()


def add_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(31, 77, 120)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    set_bidi(p)
    add_inline_runs(p, text)
    return p


def add_code_block(doc, lines):
    text = "\n".join(lines)
    p = doc.add_paragraph()
    set_bidi(p)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(31, 77, 120)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 18, "0B2545"),
        ("Heading 2", 14, "1F4D78"),
        ("Heading 3", 12, "2E74B5"),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def md_to_docx(md_path, docx_path):
    doc = Document()
    configure_document(doc)

    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_lines = []

    while idx < len(lines):
        line = lines[idx]

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if not line.strip():
            idx += 1
            continue

        if line.strip().startswith("|"):
            table_rows = []
            while idx < len(lines) and lines[idx].strip().startswith("|"):
                if not is_separator_row(lines[idx]):
                    table_rows.append(split_table_row(lines[idx]))
                idx += 1
            add_table(doc, table_rows)
            continue

        if line.startswith("# "):
            p = add_paragraph(doc, line[2:].strip())
            p.style = doc.styles["Heading 1"]
            idx += 1
            continue

        if line.startswith("## "):
            p = add_paragraph(doc, line[3:].strip())
            p.style = doc.styles["Heading 2"]
            idx += 1
            continue

        if line.startswith("### "):
            p = add_paragraph(doc, line[4:].strip())
            p.style = doc.styles["Heading 3"]
            idx += 1
            continue

        if line.startswith("- "):
            p = add_paragraph(doc, line[2:].strip(), style="List Bullet")
            set_bidi(p)
            idx += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", line)
        if numbered:
            p = add_paragraph(doc, numbered.group(2), style="List Number")
            set_bidi(p)
            idx += 1
            continue

        add_paragraph(doc, line.strip())
        idx += 1

    doc.save(docx_path)


def main():
    for md_name, docx_name in DOCS:
        md_to_docx(md_name, docx_name)
        print(docx_name)


if __name__ == "__main__":
    main()
